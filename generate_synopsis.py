"""
Generate JDMatcher Synopsis Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a1a2e')
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Light gray background via shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'f5f5f5')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)
    return p

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JDMatcher')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI-Powered Resume and Job Description Matching System')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

synopsis_label = doc.add_paragraph()
synopsis_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = synopsis_label.add_run('PROJECT SYNOPSIS')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
doc.add_paragraph()

details = [
    ('Developed by:', 'RAJ KISHOR'),
    ('Date:', 'February 2026'),
    ('Institution:', '[Your Institution Name]'),
    ('Course:', '[Your Course Name]'),
    ('Guide:', "[Your Guide's Name]"),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(value)
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════

add_heading_styled('Table of Contents', level=1)

toc_items = [
    '1. Project Title',
    '2. Project Overview',
    '3. Problem Statement',
    '4. Objectives',
    '5. Scope of the Project',
    '6. Technology Stack',
    '7. System Design',
    '8. ML Algorithm — Detailed Explanation',
    '9. Data Flow Diagrams',
    '10. User Interface Design',
    '11. Security Measures',
    '12. Testing Strategy',
    '13. Limitations',
    '14. Future Enhancements',
    '15. Conclusion',
    '16. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PROJECT TITLE
# ═══════════════════════════════════════════════════════════

add_heading_styled('1. Project Title', level=1)
p = doc.add_paragraph()
run = p.add_run('JDMatcher — AI-Powered Resume and Job Description Matching System')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════

add_heading_styled('2. Project Overview', level=1)

doc.add_paragraph(
    'JDMatcher is an intelligent, full-stack web application that leverages Machine Learning '
    'and Natural Language Processing (NLP) to automate the process of matching candidate resumes '
    'against job descriptions. The system eliminates manual resume screening by providing '
    'quantitative match scores, skill gap analysis, AI-powered recommendations, and bulk '
    'candidate ranking capabilities.'
)
doc.add_paragraph(
    'The platform serves three distinct user roles — Candidates, Recruiters, and Administrators '
    '— each with tailored dashboards, profiles, and functionalities.'
)

# ═══════════════════════════════════════════════════════════
# 3. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════

add_heading_styled('3. Problem Statement', level=1)

doc.add_paragraph('The traditional recruitment process faces several critical challenges:')
doc.add_paragraph()

problems = [
    ('Manual Resume Screening is Time-Consuming — ',
     'Recruiters spend an average of 6-7 seconds per resume, leading to superficial evaluations. '
     'For a single job posting that receives 250+ applications, this translates to hours of repetitive work.'),
    ('Inconsistent Evaluation Criteria — ',
     'Different recruiters may evaluate the same resume differently based on personal biases, '
     'fatigue, or subjective interpretation.'),
    ('Skill Gap Blindness — ',
     'Candidates often don\'t know which specific skills they lack for a particular role, '
     'leading to wasted applications and missed opportunities for upskilling.'),
    ('Scalability Issues — ',
     'When recruiters receive hundreds of resumes for a single position, it becomes practically '
     'impossible to give each resume fair consideration.'),
    ('Lack of Actionable Feedback — ',
     'Candidates receive generic rejection emails with no insight into why they weren\'t selected '
     'or what they could improve.'),
]
for i, (title, desc) in enumerate(problems, 1):
    add_bullet(desc, bold_prefix=f'{title}')

doc.add_paragraph()
doc.add_paragraph(
    'JDMatcher addresses all of these problems by automating the matching process using AI, '
    'providing transparent scoring, and generating personalized recommendations for both '
    'candidates and recruiters.'
)

# ═══════════════════════════════════════════════════════════
# 4. OBJECTIVES
# ═══════════════════════════════════════════════════════════

add_heading_styled('4. Objectives', level=1)

objectives = [
    ('1', 'Automated Resume Parsing', 'Extract structured information (skills, education, experience, projects) from unstructured PDF/DOCX resume files using NLP techniques'),
    ('2', 'Intelligent JD Parsing', 'Extract required skills, qualifications, experience levels, and responsibilities from job descriptions'),
    ('3', 'ML-Based Matching', 'Use TF-IDF vectorization and cosine similarity to compute a quantitative match score between resumes and JDs'),
    ('4', 'Weighted Multi-Factor Scoring', 'Calculate match scores using weighted components — Skills (50%), Experience (30%), Education (20%)'),
    ('5', 'Skill Gap Analysis', 'Identify missing skills and generate personalized learning recommendations with resource links'),
    ('6', 'Bulk Resume Processing', 'Allow recruiters to upload multiple resumes (including ZIP archives) and rank all candidates against a single JD'),
    ('7', 'AI-Powered Recommendations', 'Generate personalized explanations for why each candidate is or isn\'t a good fit, with comparison analysis'),
    ('8', 'Role-Based Access Control', 'Provide distinct interfaces and permissions for Candidates, Recruiters, and Administrators'),
    ('9', 'Secure Authentication', 'Implement JWT-based stateless authentication with bcrypt password hashing'),
    ('10', 'Responsive UI with Theming', 'Build a modern, mobile-responsive interface with light/dark theme support'),
]

add_table(['#', 'Objective', 'Description'], objectives, col_widths=[0.4, 1.8, 4.3])

# ═══════════════════════════════════════════════════════════
# 5. SCOPE
# ═══════════════════════════════════════════════════════════

add_heading_styled('5. Scope of the Project', level=1)

add_heading_styled('5.1 In Scope', level=2)

in_scope = [
    'User registration and authentication (JWT + bcrypt)',
    'Resume upload and parsing (PDF and DOCX formats)',
    'Job description upload and text input',
    'AI/ML-based resume-JD matching with percentage scores',
    'Skill gap analysis with learning recommendations',
    'Bulk resume upload (individual files + ZIP archives)',
    'AI-ranked candidate list with personalized explanations',
    'Role-based dashboards (Candidate, Recruiter, Admin)',
    'User profile management with language preferences',
    'Dark/Light theme toggle with persistence',
    'Contextual tips popup system for candidates and recruiters',
    'Admin panel with system statistics and user management',
    'RESTful API architecture',
    'MongoDB database for persistent storage',
]
for item in in_scope:
    add_bullet(item)

add_heading_styled('5.2 Out of Scope', level=2)

out_scope = [
    'Real-time chat or messaging between candidates and recruiters',
    'Payment/subscription system',
    'Integration with external job boards (LinkedIn, Indeed, etc.)',
    'Video interview scheduling',
    'Mobile native applications (iOS/Android)',
    'Resume building/creation tool',
]
for item in out_scope:
    add_bullet(item)

# ═══════════════════════════════════════════════════════════
# 6. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════

add_heading_styled('6. Technology Stack', level=1)

add_heading_styled('6.1 Backend', level=2)

backend_tech = [
    ('Programming Language', 'Python', '3.10+', 'Core backend logic'),
    ('Web Framework', 'Flask', '3.x', 'REST API server'),
    ('Database', 'MongoDB', '7.x', 'Document-oriented NoSQL storage'),
    ('Database Driver', 'PyMongo', '4.x', 'Python-MongoDB connector'),
    ('Authentication', 'PyJWT', '2.x', 'JSON Web Token generation/validation'),
    ('Password Hashing', 'bcrypt', '4.x', 'Secure password storage'),
    ('NLP Engine', 'spaCy', '3.x', 'Named Entity Recognition, tokenization'),
    ('Text Processing', 'NLTK', '3.x', 'Stopword removal, lemmatization'),
    ('ML Library', 'scikit-learn', '1.x', 'TF-IDF vectorizer, cosine similarity'),
    ('PDF Parser', 'PyPDF2', '3.x', 'Extract text from PDF files'),
    ('DOCX Parser', 'python-docx', '1.x', 'Extract text from Word documents'),
    ('File Handling', 'Werkzeug', '3.x', 'Secure file upload handling'),
    ('CORS', 'Flask-CORS', '4.x', 'Cross-origin request support'),
    ('ZIP Processing', 'zipfile (stdlib)', '—', 'Extract resumes from ZIP archives'),
]

add_table(['Component', 'Technology', 'Version', 'Purpose'], backend_tech, col_widths=[1.3, 1.3, 0.7, 3.2])

add_heading_styled('6.2 Frontend', level=2)

frontend_tech = [
    ('UI Library', 'React', '18.x', 'Component-based UI framework'),
    ('Language', 'TypeScript', '5.x', 'Type-safe JavaScript'),
    ('Build Tool', 'Vite', '5.x', 'Fast development server and bundler'),
    ('CSS Framework', 'Tailwind CSS', '3.x', 'Utility-first styling'),
    ('UI Components', 'shadcn/ui', 'Latest', 'Pre-built accessible React components'),
    ('Icons', 'Lucide React', 'Latest', 'Modern SVG icon library'),
    ('HTTP Client', 'Axios', '1.x', 'API communication'),
    ('Routing', 'React Router DOM', '6.x', 'Client-side routing'),
    ('Notifications', 'Sonner', '1.x', 'Toast notification system'),
    ('Theming', 'Custom Context API', '—', 'Dark/light mode with localStorage persistence'),
]

add_table(['Component', 'Technology', 'Version', 'Purpose'], frontend_tech, col_widths=[1.3, 1.3, 0.7, 3.2])

add_heading_styled('6.3 Architecture Diagram', level=2)

arch_text = """┌─────────────────────────────────────────────────────────┐
│                    CLIENT (Browser)                      │
│  ┌───────────────────────────────────────────────────┐  │
│  │           React + TypeScript + Tailwind           │  │
│  │  ┌──────────┐ ┌──────────┐ ┌──────────────────┐  │  │
│  │  │ Candidate│ │ Recruiter│ │      Admin       │  │  │
│  │  │Dashboard │ │Dashboard │ │      Panel       │  │  │
│  │  └────┬─────┘ └────┬─────┘ └───────┬──────────┘  │  │
│  │       │             │               │             │  │
│  │  ┌────┴─────────────┴───────────────┴──────────┐  │  │
│  │  │            Axios HTTP Client                │  │  │
│  │  └─────────────────────┬───────────────────────┘  │  │
│  └────────────────────────┼──────────────────────────┘  │
└───────────────────────────┼─────────────────────────────┘
                            │ REST API (JSON)
                            │ Port 8080 → Proxy → Port 5000
┌───────────────────────────┼─────────────────────────────┐
│                    SERVER (Flask)                        │
│  ┌────────────────────────┼──────────────────────────┐  │
│  │              Flask Application                    │  │
│  │  ┌──────────┐ ┌───────┴──┐ ┌──────────────────┐  │  │
│  │  │   Auth   │ │  Resume  │ │     Match        │  │  │
│  │  │  Routes  │ │  Routes  │ │     Routes       │  │  │
│  │  └────┬─────┘ └────┬─────┘ └───────┬──────────┘  │  │
│  │       │             │               │             │  │
│  │  ┌────┴─────────────┴───────────────┴──────────┐  │  │
│  │  │              Service Layer                  │  │  │
│  │  └─────────────────────┬───────────────────────┘  │  │
│  │                        │                          │  │
│  │  ┌─────────────────────┼───────────────────────┐  │  │
│  │  │           ML Engine Layer                   │  │  │
│  │  │  ┌──────────┐ ┌────┴─────┐ ┌────────────┐  │  │  │
│  │  │  │  Resume  │ │ Matching │ │  Skill Gap │  │  │  │
│  │  │  │  Parser  │ │  Engine  │ │  Analyzer  │  │  │  │
│  │  │  └──────────┘ └──────────┘ └────────────┘  │  │  │
│  │  └─────────────────────────────────────────────┘  │  │
│  └───────────────────────┬───────────────────────────┘  │
└──────────────────────────┼──────────────────────────────┘
                           │
                ┌──────────┴──────────┐
                │      MongoDB        │
                │  ┌──────────────┐   │
                │  │    users     │   │
                │  │   resumes    │   │
                │  │    jobs      │   │
                │  │   matches    │   │
                │  └──────────────┘   │
                └─────────────────────┘"""
add_code_block(arch_text)

# ═══════════════════════════════════════════════════════════
# 7. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════

add_heading_styled('7. System Design', level=1)
add_heading_styled('7.1 Module Description', level=2)

# Module 1
add_heading_styled('Module 1: Authentication & User Management', level=3)
mod1 = [
    'User registration with email validation',
    'Login with JWT token issuance (24-hour expiry)',
    'Password hashing using bcrypt (salt rounds = 12)',
    'Role-based access control (Candidate, Recruiter, Admin)',
    'Profile management (edit name, email, phone, bio, social links)',
    'Language preference selection (10 languages supported)',
    'Demo user seeding on first server start',
]
for item in mod1:
    add_bullet(item)

# Module 2
add_heading_styled('Module 2: Resume Processing', level=3)
mod2 = [
    'File upload handling (PDF, DOCX formats, max 16MB per file)',
    'Text extraction using PyPDF2 (PDF) and python-docx (DOCX)',
    'NLP-based section identification:',
]
for item in mod2:
    add_bullet(item)

sub_items = [
    'Skills Extraction — Pattern matching + NER for technical and soft skills',
    'Education Parsing — Degree, institution, year extraction',
    'Experience Parsing — Job titles, companies, duration, responsibilities',
    'Project Identification — Project names and tech stacks',
]
for item in sub_items:
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run(item)

add_bullet('Structured data storage in MongoDB')

# Module 3
add_heading_styled('Module 3: Job Description Processing', level=3)
mod3_intro = [
    'JD upload (PDF/DOCX) or direct text input',
    'Requirement extraction:',
]
for item in mod3_intro:
    add_bullet(item)

mod3_sub = [
    'Required skills and technologies',
    'Minimum experience level',
    'Education requirements',
    'Preferred qualifications',
]
for item in mod3_sub:
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run(item)

add_bullet('JD storage and management for recruiters')

# Module 4
add_heading_styled('Module 4: ML Matching Engine', level=3)
mod4 = [
    'TF-IDF Vectorization — Convert resume and JD text into numerical vectors using Term Frequency-Inverse Document Frequency',
    'Cosine Similarity — Compute angular similarity between TF-IDF vectors (range: 0.0 to 1.0)',
    'Weighted Score Calculation:',
]
for item in mod4:
    add_bullet(item)

add_code_block('Final Score = (Skills Match × 0.50) + (Experience Match × 0.30) + (Education Match × 0.20)')

mod4_extra = [
    'Skills Match — Jaccard similarity between extracted skill sets',
    'Experience Match — Comparison of years of experience vs. JD requirement',
    'Education Match — Level comparison (PhD > Masters > Bachelors > Diploma)',
]
for item in mod4_extra:
    add_bullet(item)

# Module 5
add_heading_styled('Module 5: Skill Gap Analysis', level=3)
mod5 = [
    'Identifies skills present in JD but absent from resume',
    'Categorizes missing skills by importance (Critical, Important, Nice-to-have)',
    'Generates personalized learning recommendations',
    'Suggests online resources and estimated learning timelines',
]
for item in mod5:
    add_bullet(item)

# Module 6
add_heading_styled('Module 6: Bulk Processing & AI Ranking', level=3)
mod6 = [
    'Accepts multiple resume files (PDF/DOCX) through multi-file upload',
    'ZIP archive support — automatically extracts and processes all valid resumes within',
    'Processes each resume against the provided JD',
    'Ranks all candidates by composite match score',
    'Generates per-candidate personalized AI explanations covering:',
]
for item in mod6:
    add_bullet(item)

mod6_sub = [
    'Strengths relative to JD',
    'Weaknesses and skill gaps',
    'Experience relevance assessment',
    'Education fit analysis',
]
for item in mod6_sub:
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run(item)

add_bullet('Generates top recommendation with:')
mod6_sub2 = [
    'Detailed justification for the #1 pick',
    'Comparison with the runner-up candidate',
    'Hiring readiness verdict (Ready to hire / Needs assessment / Training required)',
]
for item in mod6_sub2:
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run(item)

# Module 7
add_heading_styled('Module 7: Dashboard & Profile System', level=3)
mod7 = [
    'Candidate Dashboard: Resume summary, parsed sections display, flaws/warnings panel, match analysis interface, resume tips popup',
    'Recruiter Dashboard: Single match analysis, bulk upload zone (files + ZIP), AI-ranked candidate list, recommendation banner, detailed candidate modals',
    'Admin Panel: System statistics, user management, system settings',
    'Profile Pages (all roles): Edit personal info, language/notification preferences, role-specific settings (company details for recruiters, system config for admins)',
    'Theme Toggle: Light/dark mode switch persisted to localStorage, respects OS color scheme preference',
]
for item in mod7:
    add_bullet(item)

# 7.2 Database Design
add_heading_styled('7.2 Database Design (MongoDB Collections)', level=2)

add_heading_styled('users Collection', level=3)
users_schema = """{
  "_id": "ObjectId",
  "name": "string",
  "email": "string (unique)",
  "password": "string (bcrypt hash)",
  "role": "string (candidate | recruiter | admin)",
  "phone": "string",
  "bio": "string",
  "social_links": {
    "linkedin": "string",
    "github": "string",
    "portfolio": "string"
  },
  "preferences": {
    "language": "string",
    "email_notifications": "boolean",
    "security_alerts": "boolean"
  },
  "created_at": "datetime",
  "updated_at": "datetime"
}"""
add_code_block(users_schema)

add_heading_styled('resumes Collection', level=3)
resumes_schema = """{
  "_id": "ObjectId",
  "user_id": "ObjectId (ref: users)",
  "filename": "string",
  "raw_text": "string",
  "parsed_data": {
    "skills": ["string"],
    "education": [{
      "degree": "string",
      "institution": "string",
      "year": "number"
    }],
    "experience": [{
      "title": "string",
      "company": "string",
      "duration": "string",
      "responsibilities": ["string"]
    }],
    "projects": [{
      "name": "string",
      "technologies": ["string"],
      "description": "string"
    }]
  },
  "uploaded_at": "datetime"
}"""
add_code_block(resumes_schema)

add_heading_styled('jobs Collection', level=3)
jobs_schema = """{
  "_id": "ObjectId",
  "recruiter_id": "ObjectId (ref: users)",
  "title": "string",
  "description": "string",
  "raw_text": "string",
  "requirements": {
    "skills": ["string"],
    "experience_years": "number",
    "education_level": "string",
    "preferred": ["string"]
  },
  "created_at": "datetime"
}"""
add_code_block(jobs_schema)

add_heading_styled('matches Collection', level=3)
matches_schema = """{
  "_id": "ObjectId",
  "user_id": "ObjectId (ref: users)",
  "resume_id": "ObjectId (ref: resumes)",
  "job_id": "ObjectId (ref: jobs)",
  "scores": {
    "overall": "number (0-100)",
    "skills": "number (0-100)",
    "experience": "number (0-100)",
    "education": "number (0-100)",
    "tfidf_similarity": "number (0-1)"
  },
  "matched_skills": ["string"],
  "missing_skills": ["string"],
  "recommendations": ["string"],
  "ai_explanation": "string",
  "created_at": "datetime"
}"""
add_code_block(matches_schema)

# 7.3 API Design
add_heading_styled('7.3 API Design', level=2)

add_heading_styled('Authentication APIs', level=3)
auth_apis = [
    ('POST', '/api/auth/register', '{name, email, password, role}', '{token, user}', 'No'),
    ('POST', '/api/auth/login', '{email, password}', '{token, user}', 'No'),
    ('POST', '/api/auth/logout', '—', '{message}', 'Yes'),
    ('GET', '/api/auth/profile', '—', '{user}', 'Yes'),
]
add_table(['Method', 'Endpoint', 'Request Body', 'Response', 'Auth Required'], auth_apis, col_widths=[0.6, 1.5, 1.8, 1.3, 0.8])

add_heading_styled('Resume APIs', level=3)
resume_apis = [
    ('POST', '/api/resume/upload', 'multipart/form-data (file)', '{resume_id, parsed_data}', 'Yes'),
    ('GET', '/api/resume/list', '—', '[{resumeObjects}]', 'Yes'),
    ('GET', '/api/resume/latest', '—', '{resume}', 'Yes'),
    ('DELETE', '/api/resume/:id', '—', '{message}', 'Yes'),
]
add_table(['Method', 'Endpoint', 'Request', 'Response', 'Auth Required'], resume_apis, col_widths=[0.6, 1.5, 1.8, 1.3, 0.8])

add_heading_styled('Job Description APIs', level=3)
job_apis = [
    ('POST', '/api/job/create', '{title, description}', '{job_id}', 'Yes (Recruiter)'),
    ('POST', '/api/job/upload', 'multipart/form-data (file)', '{job_id, parsed}', 'Yes (Recruiter)'),
    ('GET', '/api/job/list', '—', '[{jobObjects}]', 'Yes (Recruiter)'),
    ('GET', '/api/job/all', '—', '[{jobObjects}]', 'Yes'),
]
add_table(['Method', 'Endpoint', 'Request', 'Response', 'Auth Required'], job_apis, col_widths=[0.6, 1.5, 1.8, 1.3, 0.8])

add_heading_styled('Matching APIs', level=3)
match_apis = [
    ('POST', '/api/match/direct', 'multipart (resume + jd)', '{scores, skills, gaps, recs}', 'Yes'),
    ('POST', '/api/match/bulk', 'multipart (resumes[] + zip + jd)', '{ranked_candidates[], rec}', 'Yes (Recruiter)'),
    ('POST', '/api/match/analyze', '{resume_id, job_id}', '{scores, analysis}', 'Yes'),
    ('GET', '/api/match/history', '—', '[{matchObjects}]', 'Yes'),
    ('POST', '/api/match/skillgap', '{resume_id, job_id}', '{missing_skills, recs}', 'Yes'),
]
add_table(['Method', 'Endpoint', 'Request', 'Response', 'Auth Required'], match_apis, col_widths=[0.6, 1.5, 1.8, 1.3, 0.8])

add_heading_styled('Admin APIs', level=3)
admin_apis = [
    ('GET', '/api/admin/stats', '—', '{total_users, total_resumes, ...}', 'Yes (Admin)'),
    ('GET', '/api/admin/users', '—', '[{userObjects}]', 'Yes (Admin)'),
]
add_table(['Method', 'Endpoint', 'Request', 'Response', 'Auth Required'], admin_apis, col_widths=[0.6, 1.5, 1.8, 1.3, 0.8])

# ═══════════════════════════════════════════════════════════
# 8. ML ALGORITHM
# ═══════════════════════════════════════════════════════════

add_heading_styled('8. ML Algorithm — Detailed Explanation', level=1)

add_heading_styled('8.1 Text Preprocessing Pipeline', level=2)
add_code_block('Raw Text → Lowercase → Remove Punctuation → Tokenize → Remove Stopwords → Lemmatize → Clean Text')

preprocess_steps = [
    ('Lowercasing — ', 'Convert all text to lowercase for uniform comparison'),
    ('Punctuation Removal — ', 'Strip special characters, numbers (context-dependent)'),
    ('Tokenization — ', 'Split text into individual words/tokens using spaCy'),
    ('Stopword Removal — ', 'Remove common words (the, is, at, which) using NLTK stopword corpus'),
    ('Lemmatization — ', 'Reduce words to base form (running → run, engineers → engineer) using spaCy'),
]
for prefix, desc in preprocess_steps:
    add_bullet(desc, bold_prefix=prefix)

add_heading_styled('8.2 TF-IDF (Term Frequency — Inverse Document Frequency)', level=2)

doc.add_paragraph(
    'TF-IDF converts text documents into numerical vectors where each dimension represents '
    'a word\'s importance.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Formula:')
run.bold = True
add_code_block(
    'TF-IDF(t, d, D) = TF(t, d) × IDF(t, D)\n\n'
    'Where:\n'
    '  TF(t, d)  = (Number of times term t appears in document d) / (Total terms in d)\n'
    '  IDF(t, D) = log(Total number of documents / Number of documents containing t)'
)

tfidf_explain = [
    ('TF (Term Frequency) — ', 'How often a word appears in the document. Higher = more relevant to that document.'),
    ('IDF (Inverse Document Frequency) — ', 'How rare a word is across all documents. Rare words get higher weight because they\'re more discriminating.'),
]
for prefix, desc in tfidf_explain:
    add_bullet(desc, bold_prefix=prefix)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Example:')
run.bold = True
add_bullet('The word "Python" appearing 5 times in a resume and being uncommon across general text gets a high TF-IDF score.')
add_bullet('The word "the" appearing 20 times but being extremely common gets a near-zero TF-IDF score.')

add_heading_styled('8.3 Cosine Similarity', level=2)

doc.add_paragraph(
    'After converting both resume and JD into TF-IDF vectors, cosine similarity measures '
    'the angle between them.'
)

p = doc.add_paragraph()
run = p.add_run('Formula:')
run.bold = True
add_code_block(
    'Cosine Similarity = (A · B) / (||A|| × ||B||)\n\n'
    'Where:\n'
    '  A · B    = Dot product of vectors A and B\n'
    '  ||A||    = Euclidean norm (magnitude) of vector A\n'
    '  ||B||    = Euclidean norm (magnitude) of vector B'
)

cosine_explain = [
    ('Result Range: ', '0.0 (completely different) to 1.0 (identical)'),
    ('Why Cosine? ', 'It measures orientation, not magnitude. A 2-page resume and a 10-page JD can still have high similarity if they discuss the same topics.'),
]
for prefix, desc in cosine_explain:
    add_bullet(desc, bold_prefix=prefix)

add_heading_styled('8.4 Weighted Composite Score', level=2)

add_code_block('Final Score = (Skills Score × 0.50) + (Experience Score × 0.30) + (Education Score × 0.20)')

score_rows = [
    ('Skills Score', '50%', '(Matched Skills Count / Total Required Skills) × 100 — Jaccard-style overlap between resume skills and JD required skills'),
    ('Experience Score', '30%', 'Comparison of candidate\'s total years vs. JD requirement. Full marks if ≥ required, proportional reduction if less'),
    ('Education Score', '20%', 'Level matching: PhD=100, Masters=80, Bachelors=60, Diploma=40. Score = (Candidate Level / Required Level) × 100, capped at 100'),
]
add_table(['Component', 'Weight', 'Calculation Method'], score_rows, col_widths=[1.2, 0.7, 4.6])

# ═══════════════════════════════════════════════════════════
# 9. DATA FLOW DIAGRAMS
# ═══════════════════════════════════════════════════════════

add_heading_styled('9. Data Flow Diagrams', level=1)

add_heading_styled('9.1 Level 0 — Context Diagram', level=2)
dfd0 = """                    ┌─────────────┐
   Resume + JD ───► │             │ ───► Match Score
                    │  JDMatcher  │ ───► Skill Gap Report
   Login Creds ───► │   System    │ ───► Recommendations
                    │             │ ───► Ranked Candidates
   Admin Query ───► │             │ ───► System Stats
                    └─────────────┘"""
add_code_block(dfd0)

add_heading_styled('9.2 Level 1 — Major Processes', level=2)
dfd1 = """┌──────────┐     ┌───────────────┐     ┌──────────────┐     ┌───────────────┐
│   User   │────►│ 1.0 Auth &    │────►│ 2.0 Document │────►│ 3.0 ML Match  │
│          │     │ User Mgmt     │     │ Processing   │     │ Engine        │
└──────────┘     └───────┬───────┘     └──────┬───────┘     └───────┬───────┘
                         │                     │                     │
                         ▼                     ▼                     ▼
                  ┌──────────────┐      ┌──────────────┐     ┌──────────────┐
                  │   MongoDB    │      │   MongoDB    │     │   MongoDB    │
                  │   (users)    │      │  (resumes,   │     │  (matches)   │
                  │              │      │   jobs)      │     │              │
                  └──────────────┘      └──────────────┘     └──────────────┘"""
add_code_block(dfd1)

add_heading_styled('9.3 Candidate Matching Flow', level=2)
candidate_flow = """Candidate uploads Resume (PDF/DOCX)
        │
        ▼
Text Extraction (PyPDF2 / python-docx)
        │
        ▼
NLP Parsing (spaCy + NLTK)
├── Skills: [Python, React, SQL, ...]
├── Education: [{BSc CS, MIT, 2022}]
├── Experience: [{SDE, Google, 2 yrs}]
└── Projects: [{E-commerce App, React+Node}]
        │
        ▼
Candidate enters/uploads JD
        │
        ▼
JD Parsing
├── Required Skills: [Python, AWS, Docker, ...]
├── Min Experience: 3 years
├── Education: Bachelor's CS
└── Preferred: [Kubernetes, CI/CD]
        │
        ▼
┌──────────────────────────────────┐
│         ML Matching Engine       │
│  ┌────────────────────────────┐  │
│  │ TF-IDF Vectorization      │  │
│  │ Resume Vector <-> JD Vector│  │
│  └────────────┬───────────────┘  │
│               ▼                  │
│  ┌────────────────────────────┐  │
│  │ Cosine Similarity = 0.73  │  │
│  └────────────┬───────────────┘  │
│               ▼                  │
│  ┌────────────────────────────┐  │
│  │ Weighted Score Calculation │  │
│  │ Skills:  80% × 0.50 = 40  │  │
│  │ Exp:     67% × 0.30 = 20  │  │
│  │ Edu:    100% × 0.20 = 20  │  │
│  │ TOTAL:           = 80%    │  │
│  └────────────────────────────┘  │
└──────────────────────────────────┘
        │
        ▼
┌──────────────────────────────────┐
│       Skill Gap Analysis         │
│  Matched: [Python, React, SQL]   │
│  Missing: [AWS, Docker]          │
│  Recommendations:                │
│  - "Learn AWS -> AWS Certified   │
│     Cloud Practitioner course"   │
│  - "Learn Docker -> Docker       │
│     official tutorial (2 weeks)" │
└──────────────────────────────────┘
        │
        ▼
Results displayed on Dashboard"""
add_code_block(candidate_flow)

add_heading_styled('9.4 Recruiter Bulk Processing Flow', level=2)
recruiter_flow = """Recruiter uploads JD + Multiple Resumes (or ZIP)
        │
        ▼
┌─────────────────────────────┐
│ If ZIP file detected:       │
│ Extract all PDF/DOCX files  │
│ from archive                │
└────────────┬────────────────┘
             ▼
For each resume file:
├── Extract text
├── Parse with NLP
├── Run ML matching against JD
├── Calculate composite score
└── Generate AI explanation
             │
             ▼
Sort all candidates by score (descending)
             │
             ▼
┌─────────────────────────────────────┐
│ AI Recommendation Generation        │
│                                     │
│ Top Pick: Candidate A (92%)         │
│ Reason: "Candidate A has 4/5       │
│ required skills including Docker    │
│ and AWS which are critical for      │
│ this DevOps role. 5 years of        │
│ experience exceeds the 3-year       │
│ requirement. Masters in CS adds     │
│ strong theoretical foundation.      │
│ Compared to runner-up Candidate B   │
│ (85%), A has cloud infrastructure   │
│ experience which is the primary     │
│ gap in B's profile."                │
│                                     │
│ Verdict: Ready to hire              │
└─────────────────────────────────────┘
             │
             ▼
Ranked list displayed with expandable
AI explanations for each candidate"""
add_code_block(recruiter_flow)

# ═══════════════════════════════════════════════════════════
# 10. USER INTERFACE DESIGN
# ═══════════════════════════════════════════════════════════

add_heading_styled('10. User Interface Design', level=1)

add_heading_styled('10.1 Page Structure', level=2)
pages = [
    ('Landing', '/', 'Public', 'Hero section, feature highlights, CTA'),
    ('Auth', '/auth', 'Public', 'Login/Register toggle form'),
    ('Candidate Dashboard', '/candidate/dashboard', 'Candidate', 'Resume upload, JD input, match analysis, resume summary, flaws panel, tips popup'),
    ('Candidate Profile', '/candidate/profile', 'Candidate', 'Edit profile, preferences (language, notifications), resume info tabs'),
    ('Recruiter Dashboard', '/recruiter/dashboard', 'Recruiter', 'Single match, bulk upload (files + ZIP), AI-ranked candidates, recommendation banner, detail modals'),
    ('Recruiter Profile', '/recruiter/profile', 'Recruiter', 'Edit profile, company details, preferences tabs, recruiting tips popup'),
    ('Admin Panel', '/admin/dashboard', 'Admin', 'System stats, user management'),
    ('Admin Profile', '/admin/profile', 'Admin', 'Edit profile, notification preferences, system settings'),
]
add_table(['Page', 'Route', 'Role', 'Key Features'], pages, col_widths=[1.2, 1.4, 0.8, 3.1])

add_heading_styled('10.2 Theme System', level=2)
theme_items = [
    ('Light Mode: ', 'Clean white backgrounds, dark text, blue accents'),
    ('Dark Mode: ', 'Dark gray/slate backgrounds, light text, blue accents'),
    ('Toggle: ', 'Sun/Moon icon button in navigation bar'),
    ('Persistence: ', 'Theme choice saved to localStorage, restored on page load'),
    ('OS Respect: ', 'On first visit, detects OS-level dark mode preference via prefers-color-scheme media query'),
]
for prefix, desc in theme_items:
    add_bullet(desc, bold_prefix=prefix)

add_heading_styled('10.3 Responsive Design', level=2)
responsive = [
    ('Desktop (≥1024px): ', 'Full sidebar navigation, multi-column layouts, detailed tables'),
    ('Tablet (768–1023px): ', 'Collapsible navigation, 2-column grids'),
    ('Mobile (<768px): ', 'Hamburger menu, single-column layouts, stacked cards'),
]
for prefix, desc in responsive:
    add_bullet(desc, bold_prefix=prefix)

# ═══════════════════════════════════════════════════════════
# 11. SECURITY MEASURES
# ═══════════════════════════════════════════════════════════

add_heading_styled('11. Security Measures', level=1)

security_rows = [
    ('Password Storage', 'bcrypt hashing with 12 salt rounds — passwords never stored in plaintext'),
    ('Authentication', 'JWT tokens with 24-hour expiry, sent via Authorization header'),
    ('Input Validation', 'Server-side validation on all endpoints (email format, password strength, file types)'),
    ('File Upload Security', 'Whitelist of allowed extensions (PDF, DOCX, ZIP only), Werkzeug secure_filename(), max size limits'),
    ('CORS', 'Configured to allow only the frontend origin'),
    ('Role-Based Access', 'JWT middleware checks user role before granting access to protected endpoints'),
    ('NoSQL Injection Prevention', 'PyMongo\'s parameterized queries prevent injection attacks'),
    ('XSS Prevention', 'React\'s built-in JSX escaping prevents cross-site scripting'),
]
add_table(['Aspect', 'Implementation'], security_rows, col_widths=[1.5, 5.0])

# ═══════════════════════════════════════════════════════════
# 12. TESTING STRATEGY
# ═══════════════════════════════════════════════════════════

add_heading_styled('12. Testing Strategy', level=1)

test_types = [
    ('Unit Testing', 'Individual functions (parsers, scoring, auth)', 'pytest, Jest'),
    ('Integration Testing', 'API endpoint behavior with database', 'pytest + Flask test client'),
    ('Frontend Testing', 'Component rendering and interaction', 'React Testing Library'),
    ('Manual Testing', 'End-to-end user flows', 'Browser-based testing'),
    ('Performance Testing', 'Bulk upload with 50+ resumes', 'Manual load testing'),
]
add_table(['Test Type', 'Scope', 'Tools'], test_types, col_widths=[1.5, 2.5, 2.5])

add_heading_styled('Test Cases', level=2)
test_cases = [
    ('1', 'Register with valid email + password', 'User created, JWT returned'),
    ('2', 'Login with wrong password', '401 Unauthorized'),
    ('3', 'Upload valid PDF resume', 'Text extracted, skills parsed'),
    ('4', 'Upload non-PDF/DOCX file', '400 Bad Request'),
    ('5', 'Match resume with JD', 'Score 0-100 returned with breakdown'),
    ('6', 'Bulk upload 10 PDFs + 1 JD', '10 ranked candidates returned'),
    ('7', 'Upload ZIP with 5 resumes', 'All 5 extracted and processed'),
    ('8', 'Access recruiter page as candidate', 'Redirect to candidate dashboard'),
    ('9', 'Theme toggle', 'UI switches dark/light, persists on reload'),
    ('10', 'Skill gap analysis', 'Missing skills listed with recommendations'),
]
add_table(['#', 'Test Case', 'Expected Result'], test_cases, col_widths=[0.4, 3.0, 3.1])

# ═══════════════════════════════════════════════════════════
# 13. LIMITATIONS
# ═══════════════════════════════════════════════════════════

add_heading_styled('13. Limitations', level=1)

limitations = [
    ('Language Support — ', 'NLP parsing is optimized for English-language resumes and JDs only. Non-English documents will produce inaccurate results.'),
    ('Resume Format Dependency — ', 'Heavily formatted resumes (tables, multi-column layouts, images) may not parse correctly from PDF.'),
    ('No Deep Learning — ', 'Uses traditional ML (TF-IDF + cosine similarity) rather than transformer models (BERT, GPT). Semantic understanding is limited.'),
    ('Skills Database — ', 'Skill extraction relies on pattern matching and a predefined skill list. Novel or niche skills may not be detected.'),
    ('No Real AI Model — ', 'The "AI explanations" are rule-based templates populated with match data, not generated by an LLM.'),
    ('Single-Server Architecture — ', 'Not designed for high-concurrency production deployment. No load balancing or horizontal scaling.'),
    ('No File Versioning — ', 'Uploading a new resume replaces the old one. No history of resume changes is maintained.'),
]
for i, (prefix, desc) in enumerate(limitations, 1):
    add_bullet(desc, bold_prefix=f'{i}. {prefix}')

# ═══════════════════════════════════════════════════════════
# 14. FUTURE ENHANCEMENTS
# ═══════════════════════════════════════════════════════════

add_heading_styled('14. Future Enhancements', level=1)

enhancements = [
    ('1', 'BERT/Transformer Integration', 'Replace TF-IDF with sentence-transformers for deeper semantic understanding'),
    ('2', 'LLM-Powered Explanations', 'Integrate OpenAI/Gemini API for truly AI-generated personalized explanations'),
    ('3', 'ATS Score Prediction', 'Predict how major ATS systems (Workday, Greenhouse) would score the resume'),
    ('4', 'Resume Builder', 'Built-in resume editor with JD-optimized suggestions'),
    ('5', 'Job Board Integration', 'Pull live JDs from LinkedIn, Indeed, Glassdoor APIs'),
    ('6', 'Interview Prep Module', 'Generate likely interview questions based on skill gaps'),
    ('7', 'Multi-Language NLP', 'Support for Hindi, Spanish, French, German resume parsing'),
    ('8', 'Email Notifications', 'Send match results and tips via email'),
    ('9', 'Docker Deployment', 'Containerize with Docker Compose for one-command deployment'),
    ('10', 'Analytics Dashboard', 'Track application trends, score distributions, popular skills over time'),
]
add_table(['#', 'Enhancement', 'Description'], enhancements, col_widths=[0.4, 2.0, 4.1])

# ═══════════════════════════════════════════════════════════
# 15. CONCLUSION
# ═══════════════════════════════════════════════════════════

add_heading_styled('15. Conclusion', level=1)

doc.add_paragraph(
    'JDMatcher demonstrates a practical application of Machine Learning and Natural Language '
    'Processing in solving a real-world recruitment problem. By automating resume screening with '
    'TF-IDF vectorization, cosine similarity scoring, and weighted multi-factor analysis, the '
    'system reduces recruiter workload while providing candidates with actionable feedback on '
    'their applications.'
)
doc.add_paragraph(
    'The three-role architecture (Candidate, Recruiter, Admin) with distinct dashboards, profile '
    'systems, and AI-powered recommendations creates a comprehensive platform that addresses the '
    'needs of all stakeholders in the hiring process. The bulk processing capability with ZIP '
    'support and per-candidate AI explanations particularly enhances the recruiter experience, '
    'enabling efficient evaluation of large applicant pools.'
)
doc.add_paragraph(
    'While the current implementation uses traditional ML techniques, the modular architecture is '
    'designed to accommodate future integration of deep learning models and LLM-based analysis, '
    'positioning the system for evolution alongside advancing AI capabilities.'
)

# ═══════════════════════════════════════════════════════════
# 16. REFERENCES
# ═══════════════════════════════════════════════════════════

add_heading_styled('16. References', level=1)

references = [
    'Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management.',
    'Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.',
    'scikit-learn Documentation — TfidfVectorizer. https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html',
    'spaCy Documentation. https://spacy.io/usage',
    'Flask Documentation. https://flask.palletsprojects.com/',
    'React Documentation. https://react.dev/',
    'MongoDB Documentation. https://www.mongodb.com/docs/',
    'NLTK Documentation. https://www.nltk.org/',
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] ')
    run.bold = True
    p.add_run(ref)

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = r'c:\Users\hp\OneDrive\Documents\jdmatcher\JDMatcher_Synopsis.docx'
doc.save(output_path)
print(f'Synopsis saved to: {output_path}')
